"""
랭킹된 논문 JSON -> Markdown 표 + Word(.docx) 레포트.

Usage:
    python build_report.py --input ranked.json --keyword "graph attention temporal" \
        --required "CVPR,SIGGRAPH" --preferred "NeurIPS,ICML,ICLR" \
        --out-md report.md --out-docx report.docx

docx 변환은 python-docx가 있으면 사용, 없으면 md만 쓰고 사용자에게 안내.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path


def truncate(text: str, n: int) -> str:
    text = (text or "").strip().replace("\n", " ").replace("\r", " ")
    text = " ".join(text.split())
    if len(text) <= n:
        return text
    return text[: n - 1].rstrip() + "…"


def md_escape(s: str) -> str:
    return (s or "").replace("|", "\\|").replace("\n", " ")


def primary_link(item: dict) -> str:
    """가장 안정적인 URL: DOI > S2 url > arXiv abs > pdf."""
    if item.get("doi"):
        return f"https://doi.org/{item['doi']}"
    if item.get("abs_url"):
        return item["abs_url"]
    if item.get("arxiv_id"):
        return f"https://arxiv.org/abs/{item['arxiv_id']}"
    if item.get("pdf_url"):
        return item["pdf_url"]
    return ""


def authors_short(authors: list, max_n: int = 3) -> str:
    if not authors:
        return ""
    if len(authors) <= max_n:
        return ", ".join(authors)
    return ", ".join(authors[:max_n]) + " et al."


def venue_short(item: dict) -> str:
    v = item.get("venue_raw") or ""
    yr = item.get("year") or ""
    if v and yr:
        return f"{v} ({yr})"
    return v or yr or "—"


def build_markdown(items: list[dict], keyword: str, required: list[str], preferred: list[str]) -> str:
    today = date.today().isoformat()
    lines = []
    lines.append(f"# Paper Finder Report")
    lines.append("")
    lines.append(f"- **Keyword:** `{keyword}`")
    lines.append(f"- **Required venues:** {', '.join(required) if required else '(none)'}")
    lines.append(f"- **Preferred venues:** {', '.join(preferred) if preferred else '(none)'}")
    lines.append(f"- **Total papers:** {len(items)}")
    lines.append(f"- **Generated:** {today}")
    lines.append("")
    lines.append("## Results")
    lines.append("")
    lines.append("| # | Title | Authors | Venue (Year) | Cites | TL;DR | Link |")
    lines.append("|---|-------|---------|--------------|-------|-------|------|")

    for i, it in enumerate(items, 1):
        title = md_escape(truncate(it.get("title", ""), 120))
        authors = md_escape(truncate(authors_short(it.get("authors") or []), 50))
        venue = md_escape(truncate(venue_short(it), 40))
        cites = it.get("citation_count", 0) or 0
        tldr = it.get("tldr") or truncate(it.get("abstract", ""), 200)
        tldr = md_escape(truncate(tldr, 240))
        link = primary_link(it)
        link_md = f"[link]({link})" if link else "—"
        lines.append(f"| {i} | {title} | {authors} | {venue} | {cites} | {tldr} | {link_md} |")

    lines.append("")
    lines.append("## Detailed Entries")
    lines.append("")
    for i, it in enumerate(items, 1):
        lines.append(f"### {i}. {it.get('title', '').strip()}")
        lines.append("")
        if it.get("authors"):
            lines.append(f"- **Authors:** {', '.join(it['authors'])}")
        lines.append(f"- **Venue:** {venue_short(it)}")
        if it.get("doi"):
            lines.append(f"- **DOI:** [{it['doi']}](https://doi.org/{it['doi']})")
        if it.get("arxiv_id"):
            lines.append(f"- **arXiv:** [{it['arxiv_id']}](https://arxiv.org/abs/{it['arxiv_id']})")
        if it.get("pdf_url"):
            lines.append(f"- **PDF:** [{it['pdf_url']}]({it['pdf_url']})")
        if it.get("citation_count"):
            lines.append(f"- **Citations:** {it['citation_count']} (influential: {it.get('influential_citation_count', 0)})")
        if it.get("tldr"):
            lines.append(f"- **TL;DR:** {it['tldr']}")
        if it.get("abstract"):
            lines.append("")
            lines.append("**Abstract:**")
            lines.append("")
            lines.append(it["abstract"])
        reasons = it.get("_reasons") or {}
        if reasons:
            lines.append("")
            lines.append(f"- _Score: {it.get('_score', 0)} ({reasons})_")
        lines.append("")
    return "\n".join(lines)


def build_docx(items: list[dict], keyword: str, required: list[str], preferred: list[str], out_path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
    except ImportError:
        print("[docx] python-docx not installed; skipping .docx output. Install with: pip install python-docx", file=sys.stderr)
        return False

    doc = Document()
    # 페이지 여백 줄여서 표 잘 들어가게
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    doc.add_heading("Paper Finder Report", level=0)
    p = doc.add_paragraph()
    p.add_run("Keyword: ").bold = True
    p.add_run(keyword)
    p = doc.add_paragraph()
    p.add_run("Required venues: ").bold = True
    p.add_run(", ".join(required) if required else "(none)")
    p = doc.add_paragraph()
    p.add_run("Preferred venues: ").bold = True
    p.add_run(", ".join(preferred) if preferred else "(none)")
    p = doc.add_paragraph()
    p.add_run("Total papers: ").bold = True
    p.add_run(str(len(items)))
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(date.today().isoformat())

    doc.add_heading("Results", level=1)

    headers = ["#", "Title", "Authors", "Venue (Year)", "Cites", "TL;DR", "Link"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for i, it in enumerate(items, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = truncate(it.get("title", ""), 200)
        cells[2].text = truncate(authors_short(it.get("authors") or []), 80)
        cells[3].text = truncate(venue_short(it), 60)
        cells[4].text = str(it.get("citation_count", 0) or 0)
        tldr = it.get("tldr") or truncate(it.get("abstract", ""), 300)
        cells[5].text = truncate(tldr, 300)

        link = primary_link(it)
        link_cell = cells[6]
        link_cell.text = ""
        if link:
            para = link_cell.paragraphs[0]
            _add_hyperlink(para, link, "link")
        else:
            link_cell.text = "—"

        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)

    doc.add_page_break()
    doc.add_heading("Detailed Entries", level=1)
    for i, it in enumerate(items, 1):
        doc.add_heading(f"{i}. {it.get('title', '').strip()}", level=2)
        if it.get("authors"):
            p = doc.add_paragraph()
            p.add_run("Authors: ").bold = True
            p.add_run(", ".join(it["authors"]))
        p = doc.add_paragraph()
        p.add_run("Venue: ").bold = True
        p.add_run(venue_short(it))

        if it.get("doi"):
            p = doc.add_paragraph()
            p.add_run("DOI: ").bold = True
            _add_hyperlink(p, f"https://doi.org/{it['doi']}", it["doi"])
        if it.get("arxiv_id"):
            p = doc.add_paragraph()
            p.add_run("arXiv: ").bold = True
            _add_hyperlink(p, f"https://arxiv.org/abs/{it['arxiv_id']}", it["arxiv_id"])
        if it.get("pdf_url"):
            p = doc.add_paragraph()
            p.add_run("PDF: ").bold = True
            _add_hyperlink(p, it["pdf_url"], it["pdf_url"])
        if it.get("citation_count"):
            p = doc.add_paragraph()
            p.add_run("Citations: ").bold = True
            p.add_run(f"{it['citation_count']} (influential: {it.get('influential_citation_count', 0)})")
        if it.get("tldr"):
            p = doc.add_paragraph()
            p.add_run("TL;DR: ").bold = True
            p.add_run(it["tldr"])
        if it.get("abstract"):
            doc.add_paragraph("Abstract:", style="Intense Quote")
            doc.add_paragraph(it["abstract"])

    doc.save(str(out_path))
    print(f"[docx] wrote {out_path}", file=sys.stderr)
    return True


_DIRECTIONS_MARKER = "Suggested Research Directions"


def md_has_directions(report_md_path: Path) -> bool:
    if not report_md_path.exists():
        return False
    return _DIRECTIONS_MARKER in report_md_path.read_text(encoding="utf-8")


def append_directions_to_md(report_md_path: Path, directions_md_path: Path) -> bool:
    """Append a directions.md (Phase 5 output) under the existing report.md.

    Returns True if appended, False if skipped (already present)."""
    if md_has_directions(report_md_path):
        return False
    existing = report_md_path.read_text(encoding="utf-8")
    add = directions_md_path.read_text(encoding="utf-8")
    report_md_path.write_text(existing + "\n\n---\n\n" + add, encoding="utf-8")
    return True


def append_directions_to_docx(report_docx_path: Path, directions_md_path: Path) -> bool:
    """Render directions.md into the existing report.docx as a new section.

    Naive markdown renderer: handles #/##/### headings, simple GitHub-style tables,
    bullet lists, blockquotes, horizontal rules, and **bold** stripping. Hyperlinks
    inside the text are kept as raw markdown text (Phase 5 output is mostly headings + tables).
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("[docx] python-docx not installed; cannot append directions to .docx.", file=sys.stderr)
        return False

    doc = Document(str(report_docx_path))
    md = directions_md_path.read_text(encoding="utf-8")

    doc.add_page_break()

    table_rows: list[list[str]] = []
    in_table = False

    def flush_table(d, rows):
        if not rows:
            return
        # rows[0] is header; if rows[1] is a separator (---), drop it
        if len(rows) >= 2 and all(set(c.strip()) <= set("-: ") for c in rows[1]):
            header = rows[0]
            data = rows[2:]
        else:
            header = rows[0]
            data = rows[1:]
        t = d.add_table(rows=1 + len(data), cols=len(header))
        t.style = "Light Grid Accent 1"
        for i, c in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = c.strip().replace("**", "")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for r, row in enumerate(data, 1):
            for c, val in enumerate(row[: len(header)]):
                cell = t.rows[r].cells[c]
                cell.text = val.strip().replace("**", "")
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8.5)

    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("|") and line.count("|") >= 2:
            cells = [c for c in line.split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("---"):
            doc.add_paragraph("—" * 30)
        elif line.startswith("> "):
            doc.add_paragraph(line[2:].strip(), style="Intense Quote")
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("_") and line.endswith("_") and len(line) > 2:
            p = doc.add_paragraph()
            r = p.add_run(line.strip("_"))
            r.italic = True
            r.font.size = Pt(9)
        else:
            doc.add_paragraph(line.replace("**", ""))

    if in_table:
        flush_table(doc, table_rows)

    doc.save(str(report_docx_path))
    print(f"[docx] appended directions to {report_docx_path}", file=sys.stderr)
    return True


def _add_hyperlink(paragraph, url: str, text: str):
    """python-docx 내장 hyperlink 없으므로 raw XML 삽입."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=False, help="ranked.json (omit when only appending directions)")
    ap.add_argument("--keyword", required=False, default="")
    ap.add_argument("--required", default="")
    ap.add_argument("--preferred", default="")
    ap.add_argument("--out-md", required=True)
    ap.add_argument("--out-docx", default=None)
    ap.add_argument(
        "--append-directions",
        default=None,
        help="Phase 5 결과(.md)를 기존 report.md/report.docx 뒤에 append. --input 없이 단독 호출 가능.",
    )
    args = ap.parse_args()

    md_path = Path(args.out_md)
    docx_path = Path(args.out_docx) if args.out_docx else None

    # Main report build (skipped if only appending directions to an existing report)
    if args.input:
        items = json.loads(Path(args.input).read_text(encoding="utf-8"))
        required = [v.strip() for v in args.required.split(",") if v.strip()]
        preferred = [v.strip() for v in args.preferred.split(",") if v.strip()]

        md = build_markdown(items, args.keyword, required, preferred)
        md_path.write_text(md, encoding="utf-8")
        print(f"[report] wrote {md_path}", file=sys.stderr)

        if docx_path:
            ok = build_docx(items, args.keyword, required, preferred, docx_path)
            if not ok:
                print(
                    "[report] .docx skipped. Run: pip install python-docx, then re-run with --out-docx.",
                    file=sys.stderr,
                )

    # Phase 5 append
    if args.append_directions:
        dir_path = Path(args.append_directions)
        if not dir_path.exists():
            print(f"[append] directions file not found: {dir_path}", file=sys.stderr)
            return 1
        if not md_path.exists():
            print(f"[append] target report.md not found: {md_path}", file=sys.stderr)
            return 1

        already_present = md_has_directions(md_path)
        # md (idempotent)
        if append_directions_to_md(md_path, dir_path):
            print(f"[append] directions -> {md_path}", file=sys.stderr)
        else:
            print(f"[append] {md_path} already contains directions, skipping md append", file=sys.stderr)

        # docx — gated on md so the pair stays consistent (we can't reliably grep docx body)
        if docx_path and docx_path.exists():
            if already_present:
                print(f"[append] md already had directions, skipping docx append for consistency: {docx_path}", file=sys.stderr)
            else:
                append_directions_to_docx(docx_path, dir_path)
        elif docx_path:
            print(f"[append] target report.docx not found, skipping docx append: {docx_path}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
